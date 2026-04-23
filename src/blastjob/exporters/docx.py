import re
from pathlib import Path


def write(resume_md: str, out_dir: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Tighten default margins
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Pt(54)
    section.left_margin = section.right_margin = Pt(54)

    lines = resume_md.splitlines()
    for line in lines:
        line = line.rstrip()

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.size = Pt(18)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:])
        elif line.startswith("---"):
            pass  # skip YAML fence markers
        elif line == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

    dest = out_dir / "resume.docx"
    doc.save(str(dest))
    return dest


def _add_inline(paragraph, text: str) -> None:
    """Parse **bold** and *italic* inline markers into runs."""
    parts = re.split(r"(\*{1,2}[^*]+\*{1,2})", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
